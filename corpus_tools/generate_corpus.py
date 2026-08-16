"""
corpus_tools/generate_corpus.py
Generates the synthetic demo corpus: the financial archive of a fictional German
Mittelstand company, "Nordlicht Logistik GmbH".

Why synthetic, and why that is the right answer rather than a shortcut:

  1. Athena's entire pitch is that sensitive documents never leave the machine.
     Committing real financial records — anyone's — to a public portfolio repo
     would contradict the premise the project is arguing for.
  2. Eval needs exact ground truth. Every figure below is generated from a fixed
     seed, so a question like "Wie hoch war der Umsatz in Q3 2025?" has one
     provably correct answer, and retrieval metrics can be scored without a
     human labelling pass.
  3. It exercises the hard cases on purpose: German and English documents, six
     file formats, tables that matter, a formula-only spreadsheet, and one
     image-only PDF that MUST surface as a visible "scanned, not indexed"
     notice rather than silently contributing nothing.

This is a one-time developer tool, not a runtime dependency. reportlab is
required only to run it and is deliberately kept out of requirements.txt.

Run:  python corpus_tools/generate_corpus.py
"""

from __future__ import annotations

import csv
import random
from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "corpus"
SEED = 20260803

COMPANY = "Nordlicht Logistik GmbH"
CITY = "Hamburg"

# ── The canonical figures. Everything else derives from these, so the corpus is
# internally consistent and the eval golden answers are exact. ────────────────
FACTS = {
    "revenue_2024": 18_452_000,
    "revenue_2023": 16_180_000,
    "revenue_q1_2025": 4_310_000,
    "revenue_q2_2025": 4_655_000,
    "revenue_q3_2025": 4_821_000,
    "opex_2024": 15_930_000,
    "ebit_2024": 2_522_000,
    "headcount_2024": 214,
    "headcount_2023": 191,
    "largest_client_share_pct": 18,
    "fleet_vehicles": 87,
    "equity_2024": 7_240_000,
    "liabilities_2024": 9_115_000,
    "total_assets_2024": 16_355_000,
    "tax_2024": 756_600,
    "audit_firm": "Wagner & Petersen Wirtschaftsprüfung GmbH",
    "largest_client": "Kranich Handels AG",
    "iban": "DE44 5001 0517 5407 3249 31",
}


def eur(n: int | float) -> str:
    """German number formatting: 18.452.000,00 — the format the corpus must use."""
    s = f"{n:,.2f}"
    return s.replace(",", "#").replace(".", ",").replace("#", ".")


# ── PDF ──────────────────────────────────────────────────────────────────────

def _pdf(path: Path, title: str, blocks: list):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=A4, title=title)
    flow = [Paragraph(title, styles["Title"]), Spacer(1, 14)]

    for kind, payload in blocks:
        if kind == "h":
            flow.append(Spacer(1, 10))
            flow.append(Paragraph(payload, styles["Heading2"]))
        elif kind == "p":
            flow.append(Paragraph(payload, styles["BodyText"]))
            flow.append(Spacer(1, 6))
        elif kind == "table":
            t = Table(payload, hAlign="LEFT")
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dde5ee")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
            ]))
            flow.append(t)
            flow.append(Spacer(1, 10))
    doc.build(flow)
    print(f"  wrote {path.name}")


def make_annual_report_de():
    f = FACTS
    _pdf(OUT / "Geschaeftsbericht_2024.pdf", f"{COMPANY} — Geschäftsbericht 2024", [
        ("p", f"{COMPANY}, {CITY}. Jahresabschluss für das Geschäftsjahr 2024, "
              f"aufgestellt nach den Vorschriften des HGB."),
        ("h", "Lagebericht"),
        ("p", f"Die Umsatzerlöse stiegen im Geschäftsjahr 2024 auf {eur(f['revenue_2024'])} EUR "
              f"nach {eur(f['revenue_2023'])} EUR im Vorjahr. Das entspricht einem Wachstum von "
              f"{(f['revenue_2024'] / f['revenue_2023'] - 1) * 100:.1f} Prozent. "
              f"Der Personalbestand erhöhte sich von {f['headcount_2023']} auf "
              f"{f['headcount_2024']} Mitarbeitende."),
        ("p", f"Der größte Einzelkunde, die {f['largest_client']}, trug "
              f"{f['largest_client_share_pct']} Prozent zum Konzernumsatz bei. Der Fuhrpark "
              f"umfasste zum Stichtag {f['fleet_vehicles']} Fahrzeuge."),
        ("h", "Gewinn- und Verlustrechnung (verkürzt)"),
        ("table", [
            ["Position", "2024 (EUR)", "2023 (EUR)"],
            ["Umsatzerlöse", eur(f["revenue_2024"]), eur(f["revenue_2023"])],
            ["Betriebliche Aufwendungen", eur(f["opex_2024"]), eur(14_010_000)],
            ["Betriebsergebnis (EBIT)", eur(f["ebit_2024"]), eur(2_170_000)],
            ["Steuern vom Einkommen und Ertrag", eur(f["tax_2024"]), eur(651_000)],
        ]),
        ("h", "Bilanz zum 31. Dezember 2024"),
        ("table", [
            ["Position", "2024 (EUR)"],
            ["Summe Aktiva", eur(f["total_assets_2024"])],
            ["Eigenkapital", eur(f["equity_2024"])],
            ["Verbindlichkeiten", eur(f["liabilities_2024"])],
        ]),
        ("h", "Bestätigungsvermerk"),
        ("p", f"Die Prüfung wurde durch die {f['audit_firm']} durchgeführt. "
              f"Der Bestätigungsvermerk wurde uneingeschränkt erteilt."),
    ])


def make_annual_report_en():
    f = FACTS
    _pdf(OUT / "Annual_Report_2024_EN.pdf", f"{COMPANY} — Annual Report 2024 (English summary)", [
        ("p", f"{COMPANY} is a logistics operator headquartered in {CITY}, Germany. "
              f"This English summary accompanies the statutory German report."),
        ("h", "Management commentary"),
        ("p", f"Revenue for the 2024 financial year reached EUR {eur(f['revenue_2024'])}, "
              f"up from EUR {eur(f['revenue_2023'])} in 2023. Headcount grew to "
              f"{f['headcount_2024']} employees. The largest single customer accounted for "
              f"{f['largest_client_share_pct']}% of consolidated revenue."),
        ("h", "Condensed income statement"),
        ("table", [
            ["Item", "2024 (EUR)", "2023 (EUR)"],
            ["Revenue", eur(f["revenue_2024"]), eur(f["revenue_2023"])],
            ["Operating expenses", eur(f["opex_2024"]), eur(14_010_000)],
            ["Operating profit (EBIT)", eur(f["ebit_2024"]), eur(2_170_000)],
        ]),
        ("p", f"The financial statements were audited by {f['audit_firm']}, "
              f"which issued an unqualified opinion."),
    ])


def make_invoice_pdf():
    f = FACTS
    net = 128_400
    vat = round(net * 0.19)
    _pdf(OUT / "Rechnung_2025_1042.pdf", "Rechnung Nr. 2025-1042", [
        ("p", f"{COMPANY}<br/>Hafenstraße 12, 20457 {CITY}<br/>USt-IdNr.: DE811907980"),
        ("p", f"Rechnungsempfänger: {f['largest_client']}, Kranichweg 4, 28195 Bremen"),
        ("p", "Rechnungsdatum: 30.09.2025 &nbsp;&nbsp; Leistungszeitraum: Q3 2025 "
              "&nbsp;&nbsp; Zahlungsziel: 30 Tage netto"),
        ("h", "Positionen"),
        ("table", [
            ["Pos.", "Beschreibung", "Menge", "Einzelpreis (EUR)", "Betrag (EUR)"],
            ["1", "Kontraktlogistik Lagerfläche", "1", eur(74_000), eur(74_000)],
            ["2", "Transportleistungen Nordroute", "1", eur(41_200), eur(41_200)],
            ["3", "Zollabwicklung", "1", eur(13_200), eur(13_200)],
        ]),
        ("table", [
            ["Zusammenfassung", "Betrag (EUR)"],
            ["Nettobetrag", eur(net)],
            ["Umsatzsteuer 19%", eur(vat)],
            ["Rechnungsbetrag brutto", eur(net + vat)],
        ]),
        ("p", f"Zahlbar auf IBAN {f['iban']}."),
    ])


def make_scanned_pdf():
    """
    An image-only PDF: no text layer at all.

    This exists so the demo can prove Athena reports "this page is a scan and was
    NOT indexed" rather than silently treating a real document as empty. In a
    financial archive, mistaking "unreadable" for "absent" is the difference
    between a caveat and a wrong answer.
    """
    from PIL import Image, ImageDraw
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    img = Image.new("RGB", (1240, 1754), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([80, 80, 1160, 300], outline="black", width=3)
    d.text((120, 150), "SPENDENQUITTUNG 2024  (eingescannt)", fill="black")
    d.text((120, 200), "Betrag: 3.500,00 EUR", fill="black")
    for i in range(12):
        d.line([120, 380 + i * 45, 1120, 380 + i * 45], fill=(60, 60, 60), width=2)
    tmp = OUT / "_scan_tmp.png"
    img.save(tmp)

    c = canvas.Canvas(str(OUT / "Spendenquittung_2024_Scan.pdf"), pagesize=A4)
    c.drawImage(ImageReader(str(tmp)), 0, 0, width=A4[0], height=A4[1])
    c.showPage()
    c.save()
    tmp.unlink()
    print("  wrote Spendenquittung_2024_Scan.pdf (image-only, no text layer)")


# ── XLSX ─────────────────────────────────────────────────────────────────────

def make_balance_xlsx():
    from openpyxl import Workbook

    f = FACTS
    wb = Workbook()

    ws = wb.active
    ws.title = "Bilanz"
    ws.append(["Position", "2024 (EUR)", "2023 (EUR)"])
    for row in [
        ["Anlagevermögen", 9_120_000, 8_540_000],
        ["Umlaufvermögen", 6_310_000, 5_890_000],
        ["Kassenbestand", 925_000, 812_000],
        ["Summe Aktiva", f["total_assets_2024"], 15_242_000],
        ["Eigenkapital", f["equity_2024"], 6_480_000],
        ["Rückstellungen", 1_240_000, 1_105_000],
        ["Verbindlichkeiten", f["liabilities_2024"], 7_657_000],
    ]:
        ws.append(row)

    ws2 = wb.create_sheet("GuV")
    ws2.append(["Position", "2024 (EUR)", "2023 (EUR)"])
    for row in [
        ["Umsatzerlöse", f["revenue_2024"], f["revenue_2023"]],
        ["Materialaufwand", 7_450_000, 6_720_000],
        ["Personalaufwand", 6_180_000, 5_410_000],
        ["Abschreibungen", 1_120_000, 1_040_000],
        ["Sonstige betriebliche Aufwendungen", 1_180_000, 840_000],
        ["Betriebsergebnis (EBIT)", f["ebit_2024"], 2_170_000],
    ]:
        ws2.append(row)

    ws3 = wb.create_sheet("Quartalsumsatz")
    ws3.append(["Quartal", "Umsatz (EUR)", "Jahr"])
    for q, v in [
        ("Q1", f["revenue_q1_2025"]), ("Q2", f["revenue_q2_2025"]), ("Q3", f["revenue_q3_2025"]),
    ]:
        ws3.append([q, v, 2025])

    wb.save(OUT / "Bilanz_und_GuV_2024.xlsx")
    print("  wrote Bilanz_und_GuV_2024.xlsx")


def make_invoice_register_xlsx():
    from openpyxl import Workbook

    rnd = random.Random(SEED)
    wb = Workbook()
    ws = wb.active
    ws.title = "Rechnungsausgang"
    ws.append(["Rechnungsnummer", "Kunde", "Datum", "Nettobetrag (EUR)", "Status"])

    clients = [FACTS["largest_client"], "Elbe Spedition GmbH", "Baltic Freight OY",
               "Rheinmetall Logistik KG", "Nordsee Fisch AG"]
    for i in range(1, 61):
        ws.append([
            f"2025-{1000 + i}",
            rnd.choice(clients),
            f"2025-0{rnd.randint(7, 9)}-{rnd.randint(10, 28)}",
            round(rnd.uniform(4_000, 140_000), 2),
            rnd.choice(["bezahlt", "offen", "bezahlt", "bezahlt"]),
        ])
    wb.save(OUT / "Rechnungsausgang_Q3_2025.xlsx")
    print("  wrote Rechnungsausgang_Q3_2025.xlsx (60 rows)")


def make_budget_xlsx_with_formulas():
    """
    A budget sheet whose totals are FORMULAS with no cached values.

    openpyxl writes no formula cache, so data_only=True returns None for these
    cells. This file is the fixture that proves Athena notices and says so,
    instead of reporting a total it could not read as absent.
    """
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Budget2026"
    ws.append(["Kostenstelle", "Q1", "Q2", "Q3", "Q4", "Summe"])
    rows = [
        ["Fuhrpark", 410_000, 425_000, 430_000, 445_000],
        ["Personal", 1_520_000, 1_545_000, 1_560_000, 1_600_000],
        ["Lager", 300_000, 305_000, 310_000, 320_000],
        ["IT", 128_000, 130_000, 133_000, 140_000],
    ]
    for i, r in enumerate(rows, start=2):
        ws.append([*r, f"=SUM(B{i}:E{i})"])
    ws.append(["Gesamt", "=SUM(B2:B5)", "=SUM(C2:C5)", "=SUM(D2:D5)", "=SUM(E2:E5)", "=SUM(F2:F5)"])
    wb.save(OUT / "Budget_2026_Plan.xlsx")
    print("  wrote Budget_2026_Plan.xlsx (formula-only totals — deliberate)")


# ── CSV ──────────────────────────────────────────────────────────────────────

def make_ledger_csv():
    rnd = random.Random(SEED + 1)
    accounts = [
        ("4400", "Umsatzerlöse Transport"), ("4401", "Umsatzerlöse Lagerung"),
        ("6000", "Personalaufwand"), ("6200", "Kfz-Kosten"),
        ("6300", "Miete Lagerfläche"), ("6800", "Bürobedarf"),
    ]
    path = OUT / "Hauptbuch_2024.csv"
    # Semicolon-delimited, cp1252 — the format German accounting software exports.
    with open(path, "w", newline="", encoding="cp1252") as fh:
        w = csv.writer(fh, delimiter=";")
        w.writerow(["Buchungsdatum", "Konto", "Bezeichnung", "Soll (EUR)", "Haben (EUR)", "Beleg"])
        for i in range(1, 241):
            acc, name = rnd.choice(accounts)
            amount = f"{rnd.uniform(200, 48000):.2f}".replace(".", ",")
            soll, haben = (amount, "") if acc.startswith("6") else ("", amount)
            w.writerow([
                f"2024-{rnd.randint(1, 12):02d}-{rnd.randint(1, 28):02d}",
                acc, name, soll, haben, f"BEL-2024-{3000 + i}",
            ])
    print("  wrote Hauptbuch_2024.csv (240 rows, cp1252, semicolon-delimited)")


def make_bank_statement_csv():
    rnd = random.Random(SEED + 2)
    path = OUT / "Kontoauszug_Q3_2025.csv"
    with open(path, "w", newline="", encoding="utf-8") as fh:
        w = csv.writer(fh, delimiter=";")
        w.writerow(["Datum", "Verwendungszweck", "Betrag (EUR)", "Saldo (EUR)"])
        saldo = 1_240_000.0
        for _ in range(80):
            amt = round(rnd.uniform(-90_000, 120_000), 2)
            saldo += amt
            w.writerow([
                f"2025-0{rnd.randint(7, 9)}-{rnd.randint(1, 28):02d}",
                rnd.choice([
                    f"Zahlungseingang {FACTS['largest_client']}",
                    "Lastschrift Kraftstoff", "Gehaltslauf", "Miete Lagerhalle",
                    "Zahlungseingang Elbe Spedition GmbH", "Versicherungsbeitrag",
                ]),
                f"{amt:.2f}".replace(".", ","),
                f"{saldo:.2f}".replace(".", ","),
            ])
    print("  wrote Kontoauszug_Q3_2025.csv (80 rows)")


# ── DOCX ─────────────────────────────────────────────────────────────────────

def make_contract_docx():
    import docx

    d = docx.Document()
    d.add_heading("Rahmenvertrag Logistikdienstleistungen", 0)
    d.add_paragraph(
        f"zwischen {COMPANY}, Hafenstraße 12, 20457 {CITY} (nachfolgend „Auftragnehmer\") "
        f"und der {FACTS['largest_client']}, Kranichweg 4, 28195 Bremen (nachfolgend „Auftraggeber\")."
    )
    d.add_heading("§ 1 Vertragsgegenstand", level=1)
    d.add_paragraph(
        "Der Auftragnehmer erbringt Kontraktlogistik-, Transport- und Zollabwicklungsleistungen "
        "für den Auftraggeber im norddeutschen Raum."
    )
    d.add_heading("§ 2 Vergütung", level=1)
    d.add_paragraph(
        "Die Vergütung erfolgt quartalsweise nach Aufwand gemäß der folgenden Preisliste. "
        "Zahlungsziel beträgt 30 Tage netto ab Rechnungsdatum."
    )
    t = d.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Leistung", "Einheit", "Preis (EUR)"
    for leistung, einheit, preis in [
        ("Lagerfläche", "je m² / Monat", "14,50"),
        ("Transport Nordroute", "je Fahrt", "820,00"),
        ("Zollabwicklung", "je Vorgang", "165,00"),
    ]:
        row = t.add_row().cells
        row[0].text, row[1].text, row[2].text = leistung, einheit, preis
    d.add_heading("§ 3 Laufzeit", level=1)
    d.add_paragraph(
        "Der Vertrag beginnt am 01.01.2024 und läuft auf unbestimmte Zeit. "
        "Er kann mit einer Frist von sechs Monaten zum Jahresende gekündigt werden."
    )
    d.save(OUT / "Rahmenvertrag_Kranich_2024.docx")
    print("  wrote Rahmenvertrag_Kranich_2024.docx")


def make_audit_docx():
    import docx

    f = FACTS
    d = docx.Document()
    d.add_heading("Independent Auditor's Report — Financial Year 2024", 0)
    d.add_paragraph(f"To the shareholders of {COMPANY}, {CITY}.")
    d.add_heading("Opinion", level=1)
    d.add_paragraph(
        f"We have audited the financial statements of {COMPANY} for the year ended "
        f"31 December 2024. In our opinion, the financial statements give a true and fair "
        f"view of the net assets and financial position of the company. We issue an "
        f"unqualified opinion."
    )
    d.add_heading("Key audit matters", level=1)
    d.add_paragraph(
        f"Revenue recognition for contract logistics was identified as a key audit matter. "
        f"Total revenue of EUR {eur(f['revenue_2024'])} was tested on a sample basis. "
        f"Customer concentration remains elevated: the largest customer represents "
        f"{f['largest_client_share_pct']}% of revenue."
    )
    d.add_heading("Summary of tested balances", level=1)
    t = d.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "Balance", "Amount (EUR)"
    for label, val in [
        ("Total assets", f["total_assets_2024"]),
        ("Equity", f["equity_2024"]),
        ("Liabilities", f["liabilities_2024"]),
        ("Income tax expense", f["tax_2024"]),
    ]:
        r = t.add_row().cells
        r[0].text, r[1].text = label, eur(val)
    d.add_paragraph(f"{CITY}, 14 March 2025 — {f['audit_firm']}")
    d.save(OUT / "Audit_Report_2024_EN.docx")
    print("  wrote Audit_Report_2024_EN.docx")


# ── Markdown / text ──────────────────────────────────────────────────────────

def make_markdown_notes():
    f = FACTS
    (OUT / "Steuerunterlagen_2024.md").write_text(
        f"""# Steuerunterlagen {COMPANY} — Veranlagungszeitraum 2024

## Körperschaft- und Gewerbesteuer

- Steuern vom Einkommen und Ertrag: **{eur(f['tax_2024'])} EUR**
- Betriebsergebnis (EBIT) als Bemessungsgrundlage: {eur(f['ebit_2024'])} EUR
- Zuständiges Finanzamt: Finanzamt {CITY}-Mitte
- Steuernummer: 22/815/04711

## Umsatzsteuer

Die Umsatzsteuer-Voranmeldungen wurden monatlich abgegeben. Der Regelsteuersatz
von 19 Prozent wurde auf alle inländischen Logistikleistungen angewandt.
Innergemeinschaftliche Lieferungen an die Baltic Freight OY wurden steuerfrei
behandelt (§ 4 Nr. 1b UStG).

## Hinweise

Der Jahresabschluss wurde von der {f['audit_firm']} geprüft. Der
Bestätigungsvermerk wurde uneingeschränkt erteilt.
""",
        encoding="utf-8",
    )
    print("  wrote Steuerunterlagen_2024.md")

    (OUT / "README_Finanzarchiv.md").write_text(
        f"""# Finanzarchiv {COMPANY}

Dieses Verzeichnis enthält die Finanzunterlagen der {COMPANY} ({CITY}).
Alle Dokumente sind **synthetisch erzeugt** und beschreiben ein fiktives
Unternehmen — sie enthalten keinerlei personenbezogene oder reale Geschäftsdaten.

| Dokument | Format | Sprache | Inhalt |
|---|---|---|---|
| Geschaeftsbericht_2024.pdf | PDF | DE | Lagebericht, GuV, Bilanz |
| Annual_Report_2024_EN.pdf | PDF | EN | English summary |
| Rechnung_2025_1042.pdf | PDF | DE | Einzelrechnung Q3 2025 |
| Spendenquittung_2024_Scan.pdf | PDF | DE | **Nur Bild — keine Textebene** |
| Bilanz_und_GuV_2024.xlsx | XLSX | DE | Bilanz, GuV, Quartalsumsatz |
| Rechnungsausgang_Q3_2025.xlsx | XLSX | DE | 60 Ausgangsrechnungen |
| Budget_2026_Plan.xlsx | XLSX | DE | **Summen nur als Formeln** |
| Hauptbuch_2024.csv | CSV | DE | 240 Buchungen (cp1252) |
| Kontoauszug_Q3_2025.csv | CSV | DE | 80 Kontobewegungen |
| Rahmenvertrag_Kranich_2024.docx | DOCX | DE | Vertrag inkl. Preisliste |
| Audit_Report_2024_EN.docx | DOCX | EN | Auditor's report |
| Steuerunterlagen_2024.md | MD | DE | Steuerliche Kennzahlen |
""",
        encoding="utf-8",
    )
    print("  wrote README_Finanzarchiv.md")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    print(f"Generating synthetic corpus into {OUT} ...")
    make_annual_report_de()
    make_annual_report_en()
    make_invoice_pdf()
    make_scanned_pdf()
    make_balance_xlsx()
    make_invoice_register_xlsx()
    make_budget_xlsx_with_formulas()
    make_ledger_csv()
    make_bank_statement_csv()
    make_contract_docx()
    make_audit_docx()
    make_markdown_notes()
    files = sorted(p.name for p in OUT.iterdir() if p.is_file())
    print(f"\nDone — {len(files)} files:")
    for n in files:
        print(f"  {n}")


if __name__ == "__main__":
    main()
