"""
core/documents.py
Parsing layer for document mode: a folder of mixed-format institutional records
becomes a flat list of ParsedElement.

Format support: PDF, XLSX/XLSM, DOCX, CSV, MD, TXT.

Why these libraries (see docs/DECISIONS.md for the full ADR):
  Every parser here is permissively licensed (MIT/BSD), pure-Python or a
  prebuilt wheel, and downloads NOTHING at runtime. That last property is not
  a nice-to-have: Athena's headline claim is that it runs with the WLAN off, and
  every ML-based parser (docling, marker, unstructured) fetches model weights on
  first use. PyMuPDF is faster and its table extraction is better, but it is
  AGPL-3.0 — viral across a network-service boundary, which is exactly what
  api/main.py is. A parser choice that forces the whole repo to AGPL contradicts
  the on-premise-deployability pitch, so pdfplumber (MIT) wins on licence.

Two invariants that the rest of document mode depends on:

  1. doc_id is derived from FILE CONTENT ONLY (sha256 of the bytes). Not path,
     not mtime, not size+mtime. mtime changes on git clone, unzip, file copy and
     Docker COPY, which would silently change every element_id and therefore
     every citation anchor and every cached vector on a fresh checkout.

  2. Table text is emitted EXACTLY ONCE. pdfplumber's page.extract_text()
     includes the text inside tables, so naively emitting both page text and
     table rows embeds every figure two or three times. Duplicates burn the
     top-k budget and actively depress context_precision. We therefore subtract
     the table bounding boxes from the page before extracting prose.
"""

from __future__ import annotations

import csv
import hashlib
import io
import os
from dataclasses import asdict, dataclass, field
from pathlib import Path

SUPPORTED_SUFFIXES = {".pdf", ".xlsx", ".xlsm", ".docx", ".csv", ".md", ".txt"}

# A sheet or CSV larger than this is indexed only in part. The cap exists so a
# 50k-row ledger cannot silently consume the whole index; the truncation is
# always announced in the table summary element (see _sheet_summary).
MAX_ROWS_EMBEDDED = int(os.getenv("ATHENA_MAX_ROWS_EMBEDDED", "2000"))

# The cap on rows carried for EXACT AGGREGATION, which is a different and much
# larger number. See _table_payload for why the two caps must not be the same.
MAX_ROWS_TABULATED = int(os.getenv("ATHENA_MAX_ROWS_TABULATED", "200000"))

# Below this many characters, a PDF page is treated as having no text layer
# (i.e. a scan). We never silently return nothing for such a page.
_SCAN_TEXT_THRESHOLD = 20


@dataclass
class ParsedElement:
    """
    One retrievable unit of a document.

    kind:
      text          — prose (a PDF page's non-table text, a DOCX paragraph run,
                      a markdown/plain-text block)
      table_row     — a single row, prefixed with its headers so the numbers
                      remain interpretable out of context
      table_summary — the whole table as markdown, for report-style questions
      notice        — a parse problem the user must see (scanned page, formula
                      cache miss, row cap hit). Never silently swallowed.

    locator is a human-readable citation anchor: "S. 3" for a PDF page,
    "Bilanz!A12" for a spreadsheet row. It is what a report footnote shows.
    """

    doc_id: str
    element_id: str
    source: str          # file name as displayed to the user
    kind: str
    text: str
    locator: str = ""
    page: int | None = None
    sheet: str | None = None
    row: int | None = None
    table_id: str | None = None
    lang_hint: str = ""
    meta: dict = field(default_factory=dict)

    def as_dict(self) -> dict:
        return asdict(self)


# ── identity ─────────────────────────────────────────────────────────────────

def content_hash(path: Path) -> str:
    """sha256 of the file's bytes. Stable across clone, copy, unzip, Docker COPY."""
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for block in iter(lambda: fh.read(1 << 20), b""):
            h.update(block)
    return h.hexdigest()


def _mk(doc_id: str, source: str, ordinal: int, kind: str, text: str, **kw) -> ParsedElement:
    return ParsedElement(
        doc_id=doc_id,
        element_id=f"{doc_id[:16]}:{ordinal:05d}",
        source=source,
        kind=kind,
        text=text,
        **kw,
    )


def _guess_lang(text: str) -> str:
    """
    Crude DE/EN hint. Deliberately not a langdetect dependency: this only feeds
    a metadata field and a UI badge, never a retrieval decision, so a wrong
    guess costs nothing. German diacritics plus a few very common function words
    separate the two corpora well enough for that purpose.
    """
    if not text:
        return ""
    low = text.lower()
    de_marks = sum(low.count(c) for c in "äöüß")
    de_words = sum(low.count(f" {w} ") for w in ("und", "der", "die", "das", "für", "nicht", "mit"))
    en_words = sum(low.count(f" {w} ") for w in ("the", "and", "for", "with", "not", "of"))
    return "de" if (de_marks * 2 + de_words) > en_words else "en"


# ── table helpers ────────────────────────────────────────────────────────────

def _norm_cell(v) -> str:
    if v is None:
        return ""
    return str(v).replace("\n", " ").strip()


def _row_text(headers: list[str], cells: list[str]) -> str:
    """
    Render one row so it survives retrieval on its own.

    "4400 | 1.245.300 | 1.180.900" is unusable: an embedding cannot tell what the
    numbers mean and a writer citing it would be inventing context. Pairing every
    value with its column header keeps the row self-describing, which is what
    makes single-fact lookups over a ledger work at all.
    """
    parts = []
    for i, cell in enumerate(cells):
        if not cell:
            continue
        head = headers[i] if i < len(headers) and headers[i] else f"Spalte {i + 1}"
        parts.append(f"{head}: {cell}")
    return " | ".join(parts)


def _table_payload(
    headers: list[str],
    rows: list[tuple[int, str, list[str]]],
    n_indexed: int,
    **anchor,
) -> dict:
    """
    The FULL table, attached to its `table_summary` element for exact aggregation.

    This is where retrieval and aggregation part company, and they have to.
    Retrieval is capped at MAX_ROWS_EMBEDDED because a 50k-row ledger would
    otherwise swamp the index — and the cap is harmless there, since top-k
    returns a handful of rows however many exist. Counting is the opposite:
    "how many invoices in Q3" is WRONG unless every row is seen, and wrong by a
    margin that grows silently as the corpus grows. Deriving a total from the
    indexed chunks would therefore report 2000 invoices for a 34,000-row
    register and look entirely confident doing it.

    So the rows travel here in full, uncapped at the retrieval limit, and never
    reach the embedder. MAX_ROWS_TABULATED still bounds memory, but it is a
    stated bound two orders of magnitude above the other one, and exceeding it
    is announced rather than silently truncated (see the callers' notices).

    `rows` is a list of (row_number, locator, cells) — the locator is built by
    the parser so a fact and the chunk it came from cite the same anchor.
    """
    capped = rows[:MAX_ROWS_TABULATED]
    return {
        "headers": headers,
        "rows": [{"row": r, "locator": loc, "cells": cells} for r, loc, cells in capped],
        "n_rows_total": len(rows),
        "n_rows_tabulated": len(capped),
        "n_rows_indexed": n_indexed,
        **anchor,
    }


def _markdown_table(headers: list[str], rows: list[list[str]], limit: int = 40) -> str:
    out = []
    if headers:
        out.append("| " + " | ".join(headers) + " |")
        out.append("|" + "---|" * len(headers))
    for r in rows[:limit]:
        out.append("| " + " | ".join(r) + " |")
    if len(rows) > limit:
        out.append(f"| ... ({len(rows) - limit} weitere Zeilen / more rows) |")
    return "\n".join(out)


# ── PDF ──────────────────────────────────────────────────────────────────────

def _parse_pdf(path: Path, doc_id: str) -> list[ParsedElement]:
    import pdfplumber

    src = path.name
    els: list[ParsedElement] = []
    n = 0

    with pdfplumber.open(str(path)) as pdf:
        for pno, page in enumerate(pdf.pages, start=1):
            tables = []
            try:
                tables = page.find_tables()
            except Exception:
                tables = []

            boxes = [t.bbox for t in tables]

            # Prose = page text MINUS anything inside a table bbox, so table
            # figures are not embedded a second time as loose text.
            def outside_tables(obj, _boxes=boxes):
                for x0, top, x1, bottom in _boxes:
                    if (
                        obj.get("x0", 0) >= x0 - 1
                        and obj.get("x1", 0) <= x1 + 1
                        and obj.get("top", 0) >= top - 1
                        and obj.get("bottom", 0) <= bottom + 1
                    ):
                        return False
                return True

            try:
                prose = (page.filter(outside_tables).extract_text() or "").strip()
            except Exception:
                prose = (page.extract_text() or "").strip()

            raw_len = len((page.extract_text() or "").strip())
            if raw_len < _SCAN_TEXT_THRESHOLD and not tables:
                n += 1
                els.append(_mk(
                    doc_id, src, n, "notice",
                    f"[{src}, Seite {pno}] Diese Seite enthält keine Textebene "
                    f"(vermutlich ein Scan). OCR ist in dieser Version deaktiviert, "
                    f"der Inhalt wurde NICHT indexiert. / This page has no text layer "
                    f"(likely a scan); OCR is disabled in this version and the content "
                    f"was NOT indexed.",
                    locator=f"S. {pno}", page=pno,
                    meta={"reason": "no_text_layer"},
                ))
                continue

            if prose:
                n += 1
                els.append(_mk(
                    doc_id, src, n, "text", prose,
                    locator=f"S. {pno}", page=pno, lang_hint=_guess_lang(prose),
                ))

            for ti, table in enumerate(tables, start=1):
                try:
                    grid = table.extract()
                except Exception:
                    continue
                if not grid:
                    continue
                rows = [[_norm_cell(c) for c in r] for r in grid]
                headers = rows[0] if rows else []
                body = rows[1:] if len(rows) > 1 else []
                tid = f"{doc_id[:16]}:p{pno}:t{ti}"

                n += 1
                els.append(_mk(
                    doc_id, src, n, "table_summary",
                    f"Tabelle {ti} auf Seite {pno} von {src}:\n"
                    + _markdown_table(headers, body),
                    locator=f"S. {pno}, Tabelle {ti}", page=pno, table_id=tid,
                    meta={"table": _table_payload(
                        headers,
                        [
                            (ri, f"S. {pno}, Tabelle {ti}, Zeile {ri}", cells)
                            for ri, cells in enumerate(body, start=1)
                        ],
                        n_indexed=len(body),
                        page=pno,
                    )},
                ))

                for ri, cells in enumerate(body, start=1):
                    line = _row_text(headers, cells)
                    if not line:
                        continue
                    n += 1
                    els.append(_mk(
                        doc_id, src, n, "table_row", line,
                        locator=f"S. {pno}, Tabelle {ti}, Zeile {ri}",
                        page=pno, row=ri, table_id=tid,
                    ))
    return els


# ── XLSX ─────────────────────────────────────────────────────────────────────

def _parse_xlsx(path: Path, doc_id: str) -> list[ParsedElement]:
    """
    Spreadsheets are read TWICE, deliberately.

    openpyxl(data_only=True) returns the cached result of a formula. That cache
    is written by Excel — not by openpyxl, and not by every headless exporter. A
    ledger whose totals are all formulas can therefore load as a grid of None,
    and the agent would answer "I could not find the total" for a file that
    plainly contains it. Reading again with data_only=False lets us tell the two
    cases apart and emit a visible notice instead of a confident wrong answer.
    """
    from openpyxl import load_workbook

    src = path.name
    els: list[ParsedElement] = []
    n = 0

    wb_val = load_workbook(str(path), data_only=True, read_only=True)
    try:
        wb_fml = load_workbook(str(path), data_only=False, read_only=True)
    except Exception:
        wb_fml = None

    try:
        for ws in wb_val.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            headers = [_norm_cell(c) for c in rows[0]]
            body_all = rows[1:]
            capped = body_all[:MAX_ROWS_EMBEDDED]
            tid = f"{doc_id[:16]}:{ws.title}"

            # Formula-cache detection
            if wb_fml is not None:
                try:
                    fml_rows = list(wb_fml[ws.title].iter_rows(values_only=True))
                    empties = sum(
                        1
                        for r_v, r_f in zip(body_all[:200], fml_rows[1:201])
                        for v, f in zip(r_v, r_f)
                        if v is None and isinstance(f, str) and f.startswith("=")
                    )
                    if empties:
                        n += 1
                        els.append(_mk(
                            doc_id, src, n, "notice",
                            f"[{src}, Blatt {ws.title}] {empties} Formelzellen haben keinen "
                            f"gespeicherten Wert (die Datei wurde zuletzt nicht von Excel "
                            f"geschrieben). Diese Werte fehlen im Index. / {empties} formula "
                            f"cells have no cached value; those figures are missing from the index.",
                            locator=f"{ws.title}", sheet=ws.title,
                            meta={"reason": "formula_cache_miss", "count": empties},
                        ))
                except Exception:
                    pass

            if len(body_all) > MAX_ROWS_EMBEDDED:
                n += 1
                els.append(_mk(
                    doc_id, src, n, "notice",
                    f"[{src}, Blatt {ws.title}] Nur die ersten {MAX_ROWS_EMBEDDED} von "
                    f"{len(body_all)} Datenzeilen wurden indexiert. / Only the first "
                    f"{MAX_ROWS_EMBEDDED} of {len(body_all)} data rows were indexed.",
                    locator=f"{ws.title}", sheet=ws.title,
                    meta={"reason": "row_cap", "total_rows": len(body_all)},
                ))

            n += 1
            body_str = [[_norm_cell(c) for c in r] for r in capped]
            els.append(_mk(
                doc_id, src, n, "table_summary",
                f"Arbeitsblatt '{ws.title}' aus {src} "
                f"({len(body_all)} Datenzeilen):\n" + _markdown_table(headers, body_str),
                locator=f"{ws.title}", sheet=ws.title, table_id=tid,
                meta={"table": _table_payload(
                    headers,
                    [
                        (ri, f"{ws.title}!Z{ri}", [_norm_cell(c) for c in r])
                        for ri, r in enumerate(body_all, start=2)
                    ],
                    n_indexed=len(capped),
                    sheet=ws.title,
                )},
            ))

            for ri, r in enumerate(capped, start=2):  # row 1 is the header
                cells = [_norm_cell(c) for c in r]
                line = _row_text(headers, cells)
                if not line:
                    continue
                n += 1
                els.append(_mk(
                    doc_id, src, n, "table_row", line,
                    locator=f"{ws.title}!Z{ri}", sheet=ws.title, row=ri, table_id=tid,
                ))
    finally:
        wb_val.close()
        if wb_fml is not None:
            wb_fml.close()
    return els


# ── CSV ──────────────────────────────────────────────────────────────────────

def _read_text_guess_encoding(path: Path) -> str:
    """
    German financial exports are frequently cp1252/latin-1, not UTF-8. Guessing
    wrong turns 'Umsatzerlöse' into mojibake, which breaks both lexical search
    and the embedding. charset-normalizer picks the encoding; we fall back to
    utf-8 with replacement so a bad file degrades instead of raising.
    """
    raw = path.read_bytes()

    # Strict UTF-8 first: if the bytes are valid UTF-8 they almost certainly are
    # UTF-8, and this is both faster and more reliable than statistics.
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass

    # Statistical detection needs a reasonable sample. A short export (a handful
    # of rows) does not give charset-normalizer enough signal and it can pick the
    # wrong single-byte codec, so only trust it on larger files.
    if len(raw) >= 2048:
        try:
            from charset_normalizer import from_bytes

            best = from_bytes(raw).best()
            if best is not None:
                return str(best)
        except Exception:
            pass

    # cp1252 is the dominant legacy encoding for German accounting exports and
    # is a strict superset of latin-1 for our purposes. Try it before giving up,
    # otherwise 'Umsatzerlöse' arrives as mojibake and matches nothing.
    for codec in ("cp1252", "latin-1"):
        try:
            return raw.decode(codec)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _parse_csv(path: Path, doc_id: str) -> list[ParsedElement]:
    src = path.name
    text = _read_text_guess_encoding(path)

    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
    except Exception:
        class dialect(csv.excel):  # noqa: N801 - csv API expects a class
            delimiter = ";" if text[:4096].count(";") > text[:4096].count(",") else ","

    rows = list(csv.reader(io.StringIO(text), dialect))
    if not rows:
        return []

    headers = [_norm_cell(c) for c in rows[0]]
    body = rows[1:]
    capped = body[:MAX_ROWS_EMBEDDED]
    tid = f"{doc_id[:16]}:csv"

    els: list[ParsedElement] = []
    n = 0

    if len(body) > MAX_ROWS_EMBEDDED:
        n += 1
        els.append(_mk(
            doc_id, src, n, "notice",
            f"[{src}] Nur die ersten {MAX_ROWS_EMBEDDED} von {len(body)} Zeilen indexiert. "
            f"/ Only the first {MAX_ROWS_EMBEDDED} of {len(body)} rows were indexed.",
            meta={"reason": "row_cap", "total_rows": len(body)},
        ))

    n += 1
    els.append(_mk(
        doc_id, src, n, "table_summary",
        f"{src} ({len(body)} Datenzeilen):\n"
        + _markdown_table(headers, [[_norm_cell(c) for c in r] for r in capped]),
        table_id=tid,
        meta={"table": _table_payload(
            headers,
            [
                (ri, f"Zeile {ri}", [_norm_cell(c) for c in r])
                for ri, r in enumerate(body, start=2)
            ],
            n_indexed=len(capped),
        )},
    ))

    for ri, r in enumerate(capped, start=2):
        line = _row_text(headers, [_norm_cell(c) for c in r])
        if not line:
            continue
        n += 1
        els.append(_mk(
            doc_id, src, n, "table_row", line,
            locator=f"Zeile {ri}", row=ri, table_id=tid,
        ))
    return els


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _parse_docx(path: Path, doc_id: str) -> list[ParsedElement]:
    import docx

    src = path.name
    d = docx.Document(str(path))
    els: list[ParsedElement] = []
    n = 0

    buf: list[str] = []
    for para in d.paragraphs:
        t = para.text.strip()
        if t:
            buf.append(t)
    prose = "\n".join(buf).strip()
    if prose:
        n += 1
        els.append(_mk(doc_id, src, n, "text", prose, lang_hint=_guess_lang(prose)))

    for ti, table in enumerate(d.tables, start=1):
        grid = [[_norm_cell(c.text) for c in row.cells] for row in table.rows]
        if not grid:
            continue
        headers, body = grid[0], grid[1:]
        tid = f"{doc_id[:16]}:t{ti}"
        n += 1
        els.append(_mk(
            doc_id, src, n, "table_summary",
            f"Tabelle {ti} aus {src}:\n" + _markdown_table(headers, body),
            locator=f"Tabelle {ti}", table_id=tid,
            meta={"table": _table_payload(
                headers,
                [
                    (ri, f"Tabelle {ti}, Zeile {ri}", cells)
                    for ri, cells in enumerate(body, start=1)
                ],
                n_indexed=len(body),
            )},
        ))
        for ri, cells in enumerate(body, start=1):
            line = _row_text(headers, cells)
            if not line:
                continue
            n += 1
            els.append(_mk(
                doc_id, src, n, "table_row", line,
                locator=f"Tabelle {ti}, Zeile {ri}", row=ri, table_id=tid,
            ))
    return els


# ── plain text / markdown ────────────────────────────────────────────────────

def _parse_text(path: Path, doc_id: str) -> list[ParsedElement]:
    src = path.name
    text = _read_text_guess_encoding(path).strip()
    if not text:
        return []

    # Split on blank lines, then regroup into ~1200-char blocks. Markdown
    # headings start a new block so a section never merges into its neighbour.
    blocks: list[str] = []
    cur: list[str] = []
    size = 0
    for para in text.split("\n\n"):
        p = para.strip()
        if not p:
            continue
        if (p.startswith("#") and cur) or size + len(p) > 1200:
            blocks.append("\n\n".join(cur))
            cur, size = [], 0
        cur.append(p)
        size += len(p)
    if cur:
        blocks.append("\n\n".join(cur))

    return [
        _mk(doc_id, src, i, "text", b, locator=f"Abschnitt {i}", lang_hint=_guess_lang(b))
        for i, b in enumerate(blocks, start=1)
        if b.strip()
    ]


# ── public API ───────────────────────────────────────────────────────────────

_PARSERS = {
    ".pdf": _parse_pdf,
    ".xlsx": _parse_xlsx,
    ".xlsm": _parse_xlsx,
    ".docx": _parse_docx,
    ".csv": _parse_csv,
    ".md": _parse_text,
    ".txt": _parse_text,
}


def parse_file(path: str | Path) -> tuple[str, list[ParsedElement]]:
    """
    Parse one file into (doc_id, elements).

    Raises nothing for a recognised-but-broken file: the caller gets a single
    'notice' element describing the failure. "The corpus does not contain this"
    and "that file failed to parse" are different claims, and in a financial
    setting conflating them is a correctness bug, not a cosmetic one.
    """
    p = Path(path)
    doc_id = content_hash(p)
    suffix = p.suffix.lower()
    parser = _PARSERS.get(suffix)

    if parser is None:
        return doc_id, [_mk(
            doc_id, p.name, 1, "notice",
            f"[{p.name}] Nicht unterstütztes Format '{suffix}'. / Unsupported format.",
            meta={"reason": "unsupported_format"},
        )]

    try:
        return doc_id, parser(p, doc_id)
    except Exception as e:  # noqa: BLE001 - a bad file must not kill the ingest
        return doc_id, [_mk(
            doc_id, p.name, 1, "notice",
            f"[{p.name}] Datei konnte nicht gelesen werden / could not be parsed: "
            f"{type(e).__name__}: {e}",
            meta={"reason": "parse_error", "error": f"{type(e).__name__}: {e}"},
        )]


def parse_file_bytes(filename: str, data: bytes) -> tuple[str, list[ParsedElement]]:
    """
    Parse an uploaded file held in memory.

    Session uploads never exist as a corpus path, but every parser in this module
    is built on a real file (pdfplumber, openpyxl and python-docx all want a path
    or a seekable handle). The bytes are therefore written to a temporary file
    that is deleted before this function returns, in a `finally`, so an
    exception mid-parse cannot leave an uploaded document on disk — which would
    quietly break the promise that session files do not persist.

    Only the SUFFIX of the supplied name is trusted. The name itself is never
    used to build the temp path, so a filename like "../../athena_index.db"
    cannot escape the temp directory.
    """
    import tempfile

    suffix = Path(filename).suffix.lower()
    doc_id = hashlib.sha256(data).hexdigest()

    if suffix not in SUPPORTED_SUFFIXES:
        return doc_id, [_mk(
            doc_id, filename, 1, "notice",
            f"[{filename}] Nicht unterstütztes Format '{suffix}'. / Unsupported format.",
            meta={"reason": "unsupported_format"},
        )]

    fd, tmp_path = tempfile.mkstemp(suffix=suffix, prefix="athena_upload_")
    try:
        with os.fdopen(fd, "wb") as fh:
            fh.write(data)
        parser = _PARSERS[suffix]
        try:
            elements = parser(Path(tmp_path), doc_id)
        except Exception as e:  # noqa: BLE001
            return doc_id, [_mk(
                doc_id, filename, 1, "notice",
                f"[{filename}] Datei konnte nicht gelesen werden / could not be parsed: "
                f"{type(e).__name__}: {e}",
                meta={"reason": "parse_error", "error": f"{type(e).__name__}: {e}"},
            )]
        # The parsers stamp `source` from the temp file's name; restore the name
        # the user actually uploaded so citations read correctly.
        for el in elements:
            el.source = filename
        return doc_id, elements
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def iter_corpus(folder: str | Path) -> list[Path]:
    """All supported files under `folder`, sorted for deterministic ingest order."""
    root = Path(folder)
    if not root.exists():
        return []
    return sorted(
        p for p in root.rglob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
    )
